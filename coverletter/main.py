import cmd
from docx import Document
from . import separators


class CoverLetterParser(cmd.Cmd):
  
  """
    Initiates an Instance of Parser
  """
  
  intro = "Simple Cover Letter Parser" 

  def do_prompt(self, line):
    #Present a link to a word document
    
    self.directory = input('Please input the absolute path to the local document')
    file_template = open(directory, 'rb')
    self.document = Document(file_template)
    parseDocument(document)

  #Once Open Parse the Document for Separators, defaults to [];
  def parseDocument(self):
    """
      Parse the document
    """
    #Go through the document
    for paragraph in self.document.paragraphs:
      print paragraph.text
    
    

  #Reads Inputs for Specific Headers Supported are within the list  
  @classmethod
  def insert_overwrite(self, text):
    """
      Overwrite the separators with expected input;
    """
    print(text)
    new_text = input('What new text do you want to put inside?')

  @classmethod
  def save_input_field(self, answer):
    """
      For same fields, save them to a database so we can reuse the same input
    """
    if not self.database:
      self.database = {}
      self.database[inputfield] = answer
    else: 
      self.database[inputfield] = answer

  @staticmethod
  def write_document(document):
    path = input('Where would you like to save the document?')
    new_name = input('What is the filename... example [FName-LName-Date]')
    document.save(path + new_name)
    

if __name__ == '__main__':
  CoverLetterParser().cmdloop()
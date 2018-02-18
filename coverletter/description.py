from textblob import TextBlob
from docx import Document

class JobDescriptionParser:
  
  def __init__(self, filepath):
    self.parse_document(filepath)
    self.output = self.top_nouns(self.document)

  def parse_document(self, filepath):
    """
      Parses text document so that all paragraphs are accessible
    """
    with open(filepath, 'rb') as file_template:
      self.document = Document(file_template)
      
  def top_nouns(self, document):
    """
      Returns the total document paragraphs
    """
    
    text_string = " ".join(paragraph.text for paragraph in document.paragraphs)

    total_text = TextBlob(text_string)
    test_dict = total_text.np_counts
    sorted_dict = sorted(test_dict, key=test_dict.get, reverse=True)
    return sorted_dict[0:9]

  def top_words(self, document):
    """
      Return top single words written in the document
    """
    totalstring = " ".join(paragraph.text for paragraph in document.paragraphs)
    word_text = TextBlob(totalstring)
    word_dict = word_text.word_counts
    sorted_dict = sorted(word_dict, key=word_dict.get, reverse=True)
    return sorted_dict[0:9]

  

if __name__ == '__main__':
  JobDescriptionParser()

